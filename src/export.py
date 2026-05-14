"""
Export Module
==============
Handles exporting OCR results to multiple output formats.

Supported formats:
    Clipboard — instant copy via pyperclip (no file created)
    TXT       — plain text with optional metadata header
    DOCX      — formatted Word document with metadata table and footer
    JSON      — full result dict including all metadata fields
    CSV       — batch export, one row per capture session

All file exports write to the output_dir specified at construction
(default: ocr_output/) which is created automatically if needed.

Dependencies:
    pyperclip   — clipboard export (optional, gracefully disabled if absent)
    python-docx — DOCX export (optional, gracefully disabled if absent)
    All other formats use the Python standard library only.

Author: Harry Larkin
Project: Screenshot OCR Tool (CI601)
Date: January 2026
"""

import os
import json
import csv
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List

try:
    import pyperclip
    CLIPBOARD_AVAILABLE = True
except ImportError:
    CLIPBOARD_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class OCRExporter:
    """
    Unified export interface for OCR results.

    Wraps clipboard copy and file export functions, passing the shared
    output directory to each call. Provides export_all_formats() to
    run all formats in a single call.

    Usage:
        exporter = OCRExporter()
        exporter.export_to_clipboard(text)
        exporter.export_to_docx(text, metadata=result)
    """

    def __init__(self, output_dir: str = "ocr_output"):
        """
        Args:
            output_dir: Directory for all file exports. Created automatically
                        if it does not already exist.
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    # ── Clipboard ──────────────────────────────────────────────────────────────

    def export_to_clipboard(self, text: str) -> bool:
        """
        Copy text to the system clipboard via pyperclip.

        Returns:
            True if successful, False if pyperclip is unavailable or fails.
        """
        if not CLIPBOARD_AVAILABLE:
            print("Clipboard unavailable — install pyperclip: pip install pyperclip")
            return False
        try:
            pyperclip.copy(text)
            print(f"Text copied to clipboard ({len(text)} characters).")
            return True
        except Exception as e:
            print(f"Clipboard copy failed: {e}")
            return False

    # ── TXT ────────────────────────────────────────────────────────────────────

    def export_to_txt(
        self,
        text:             str,
        filepath:         Optional[str]           = None,
        include_metadata: bool                    = False,
        metadata:         Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        """
        Save extracted text as a plain text file.

        If include_metadata is True, a header block is prepended showing
        the extraction timestamp, confidence, character count, word count
        and processing time.

        Args:
            text:             Extracted text to save.
            filepath:         Output path. Auto-generated if None.
            include_metadata: Whether to prepend a metadata header.
            metadata:         OCREngine result dict for the header.

        Returns:
            Absolute path to the saved file, or None if saving failed.
        """
        if not filepath:
            ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / f"ocr_text_{ts}.txt"
        else:
            filepath = Path(filepath)

        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                if include_metadata and metadata:
                    f.write("=" * 70 + "\n")
                    f.write("OCR EXTRACTED TEXT\n")
                    f.write("=" * 70 + "\n\n")
                    f.write(f"Extraction Date : {metadata.get('timestamp', 'N/A')}\n")
                    f.write(f"Confidence      : {metadata.get('confidence', 0):.1f}%\n")
                    f.write(f"Characters      : {metadata.get('char_count', 'N/A')}\n")
                    f.write(f"Words           : {metadata.get('word_count', 'N/A')}\n")
                    f.write(f"Processing Time : {metadata.get('processing_time', 0):.2f}s\n")
                    f.write("\n" + "-" * 70 + "\n\n")
                f.write(text)

            print(f"TXT saved: {filepath}  ({filepath.stat().st_size} bytes)")
            return str(filepath)

        except Exception as e:
            print(f"TXT export failed: {e}")
            return None

    # ── DOCX ───────────────────────────────────────────────────────────────────

    def export_to_docx(
        self,
        text:             str,
        filepath:         Optional[str]           = None,
        include_metadata: bool                    = True,
        metadata:         Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        """
        Save extracted text as a formatted Word document.

        Document structure:
            Title          — 'OCR Extracted Text' (centred heading)
            Metadata table — 5-row label/value table (if enabled)
            Extracted text — one paragraph per line at Calibri 11pt
            Footer         — generation timestamp

        Requires python-docx. Returns None gracefully if not installed.

        Args:
            text:             Extracted text to save.
            filepath:         Output path. Auto-generated if None.
            include_metadata: Whether to include the metadata table.
            metadata:         OCREngine result dict.

        Returns:
            Absolute path to the saved .docx, or None if saving failed.
        """
        if not DOCX_AVAILABLE:
            print("DOCX unavailable — install python-docx: pip install python-docx")
            return None

        if not filepath:
            ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / f"ocr_document_{ts}.docx"
        else:
            filepath = Path(filepath)

        try:
            doc   = Document()
            title = doc.add_heading('OCR Extracted Text', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            if include_metadata and metadata:
                doc.add_heading('Extraction Metadata', level=1)
                table = doc.add_table(rows=5, cols=2)
                table.style = 'Light Grid Accent 1'

                for i, (label, value) in enumerate([
                    ('Extraction Date',  metadata.get('timestamp', 'N/A')),
                    ('Confidence Score', f"{metadata.get('confidence', 0):.1f}%"),
                    ('Character Count',  str(metadata.get('char_count', 'N/A'))),
                    ('Word Count',       str(metadata.get('word_count', 'N/A'))),
                    ('Processing Time',  f"{metadata.get('processing_time', 0):.2f}s"),
                ]):
                    row               = table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = value
                    row.cells[0].paragraphs[0].runs[0].bold = True

                doc.add_paragraph()

            doc.add_heading('Extracted Text', level=1)
            for line in text.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)

            footer           = doc.sections[0].footer
            footer_para      = footer.paragraphs[0]
            footer_para.text = (
                f"Generated by Screenshot OCR Tool — "
                f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.save(filepath)
            print(f"DOCX saved: {filepath}  ({filepath.stat().st_size} bytes)")
            return str(filepath)

        except Exception as e:
            print(f"DOCX export failed: {e}")
            return None

    # ── JSON ───────────────────────────────────────────────────────────────────

    def export_to_json(
        self,
        text:     str,
        filepath: Optional[str]           = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        """
        Save extracted text and metadata as a JSON file.

        JSON structure:
            { "text": "...", "metadata": {...}, "export_info": {...} }

        Args:
            text:     Extracted text to save.
            filepath: Output path. Auto-generated if None.
            metadata: OCREngine result dict.

        Returns:
            Absolute path to the saved .json, or None if saving failed.
        """
        if not filepath:
            ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / f"ocr_data_{ts}.json"
        else:
            filepath = Path(filepath)

        try:
            data = {
                'text':        text,
                'metadata':    metadata or {},
                'export_info': {
                    'export_timestamp': datetime.now().isoformat(),
                    'format_version':   '1.0',
                }
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            print(f"JSON saved: {filepath}  ({filepath.stat().st_size} bytes)")
            return str(filepath)

        except Exception as e:
            print(f"JSON export failed: {e}")
            return None

    # ── CSV ────────────────────────────────────────────────────────────────────

    def export_to_csv(
        self,
        records:  List[Dict[str, Any]],
        filepath: Optional[str] = None,
    ) -> Optional[str]:
        """
        Save a batch of OCR results as a CSV file.

        Each row represents one capture. Newlines in text are replaced
        with spaces to keep the CSV valid.

        Args:
            records:  List of OCREngine result dicts.
            filepath: Output path. Auto-generated if None.

        Returns:
            Absolute path to the saved .csv, or None if saving failed.
        """
        if not records:
            print("CSV export: no records.")
            return None

        if not filepath:
            ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = self.output_dir / f"ocr_batch_{ts}.csv"
        else:
            filepath = Path(filepath)

        fieldnames = ['timestamp', 'text', 'confidence',
                      'char_count', 'word_count', 'processing_time']

        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for r in records:
                    writer.writerow({
                        'timestamp':       r.get('timestamp', ''),
                        'text':            r.get('text', '').replace('\n', ' '),
                        'confidence':      r.get('confidence', 0),
                        'char_count':      r.get('char_count', 0),
                        'word_count':      r.get('word_count', 0),
                        'processing_time': r.get('processing_time', 0),
                    })

            print(f"CSV saved: {filepath}  ({len(records)} records)")
            return str(filepath)

        except Exception as e:
            print(f"CSV export failed: {e}")
            return None

    # ── All formats ────────────────────────────────────────────────────────────

    def export_all_formats(
        self,
        text:          str,
        metadata:      Optional[Dict[str, Any]] = None,
        base_filename: Optional[str]            = None,
    ) -> Dict[str, Optional[str]]:
        """
        Export to clipboard, TXT, JSON and DOCX simultaneously.

        Args:
            text:          Extracted text to export.
            metadata:      OCREngine result dict.
            base_filename: Shared base filename (without extension).

        Returns:
            Dict mapping format names to saved paths (None if failed).
        """
        if not base_filename:
            base_filename = f"ocr_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        results = {
            'clipboard': self.export_to_clipboard(text),
            'txt':  self.export_to_txt(
                text,
                str(self.output_dir / f"{base_filename}.txt"),
                include_metadata=True, metadata=metadata
            ),
            'json': self.export_to_json(
                text,
                str(self.output_dir / f"{base_filename}.json"),
                metadata=metadata
            ),
            'docx': self.export_to_docx(
                text,
                str(self.output_dir / f"{base_filename}.docx"),
                include_metadata=True, metadata=metadata
            ),
        }

        successful = sum(1 for v in results.values() if v)
        print(f"Export complete: {successful}/{len(results)} formats succeeded.")
        return results